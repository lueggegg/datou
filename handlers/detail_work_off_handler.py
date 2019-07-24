# -*- coding: utf-8 -*-

from base_handler import BaseHandler, HttpClient
from tornado import gen
from datetime import datetime
import type_define
from docx import Document
import os
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

job_sequence_map = {
    type_define.TYPE_JOB_ASK_FOR_LEAVE_NORMAL_IN_ONE_DAY_NEW: {
        type_define.job_sequence_add: type_define.job_sequence_pre_judge,
        type_define.job_sequence_pre_judge: type_define.job_sequence_leader_judge,
        type_define.job_sequence_leader_judge: type_define.job_sequence_hr_record,
    },
    type_define.TYPE_JOB_ASK_FOR_LEAVE_NORMAL_BEYOND_ONE_DAY_NEW: {
        type_define.job_sequence_add: type_define.job_sequence_pre_judge,
        type_define.job_sequence_pre_judge: type_define.job_sequence_leader_judge,
        type_define.job_sequence_leader_judge: type_define.job_sequence_via_leader_judge,
        type_define.job_sequence_via_leader_judge: type_define.job_sequence_hr_record,
    },
    type_define.TYPE_JOB_ASK_FOR_LEAVE_LEADER_BEYOND_ONE_DAY_NEW: {
        type_define.job_sequence_add: type_define.job_sequence_pre_judge,
        type_define.job_sequence_pre_judge: type_define.job_sequence_leader_judge,
        type_define.job_sequence_leader_judge: type_define.job_sequence_via_leader_judge,
        type_define.job_sequence_via_leader_judge: type_define.job_sequence_main_leader_judge,
        type_define.job_sequence_main_leader_judge: type_define.job_sequence_hr_record,
    }
}
job_sequence_map[type_define.TYPE_JOB_ASK_FOR_LEAVE_LEADER_IN_ONE_DAY_NEW] = job_sequence_map[type_define.TYPE_JOB_ASK_FOR_LEAVE_NORMAL_BEYOND_ONE_DAY_NEW]


job_status_desc = {}
job_status_desc[type_define.STATUS_JOB_PROCESSING] = '<span style="color: orange">处理中</span>'
job_status_desc[type_define.STATUS_JOB_COMPLETED] = '<span style="color: rgb(0,225,32)">已完成</span>'
job_status_desc[type_define.STATUS_JOB_REJECTED] = '<span style="color: red">未通过</span>'
job_status_desc[type_define.STATUS_JOB_CANCEL] = '<span style="color: gray">已撤回</span>'
job_status_desc[type_define.STATUS_JOB_SYS_CANCEL] = '<span style="color: red">系统撤回</span>'

seq_desc_map = {
    type_define.job_sequence_add: '申请休假',
    type_define.job_sequence_pre_judge: '人事部初核意见',
    type_define.job_sequence_leader_judge: '部门负责人意见',
    type_define.job_sequence_via_leader_judge: '分管领导意见',
    type_define.job_sequence_hr_leader_judge: '分管人事领导意见',
    type_define.job_sequence_main_leader_judge: '主要负责人意见',
    type_define.job_sequence_hr_record: '人事部备案情况',
}

class DetailWorkOffHandler(BaseHandler):

    @gen.coroutine
    def get(self, *args, **kwargs):
        st = yield self.verify_user()
        if not st:
            return
        job_id = int(self.get_argument_and_check_it('job_id'))
        job_record = yield self.job_dao.query_job_base_info(job_id)
        job_nodes = yield self.job_dao.query_job_node_list(job_id)
        leave_detail = yield self.job_dao.query_new_leave_detail(job_id)
        annual = yield self.get_annual(job_record['invoker'])
        job_type = job_record['type']
        is_roll = self.is_roll(job_record)
        detail = {
            'title': '请假' if not is_roll else '销假',
            'job_id': job_id,
            'invoker': '',
            'department': '',
            'type': leave_detail['leave_type'],
            'total_annual': annual['total'],
            'used_annual': annual['used'],
            'pre_judgement': '',
            'reason': '',
            'time': self.get_leave_time_desc(leave_detail),
            'department_leader_judgement': '',
            'via_leader_judgement': '',
            'hr_leader_judgement': '',
            'main_leader_judgement': '',
            'hr_department_record': '',
            'attachment': [],
            'waiting': True,
            'process_desc': '审批意见',
            'can_cancel': True,
            'can_export': False,
            'can_roll_back': False,
            'can_del': self.account_info['id'] == 250,
            'status': '处理中',
            'reply_desc': '拟同意',
            'reject_desc': '不同意',
            'mark': None,
        }
        if leave_detail['extend']:
            detail['type'] += '(%s)' % leave_detail['extend']
        if leave_detail['annual_part']:
            annual_part = leave_detail['annual_part'] / 2.0
            off_part = leave_detail['off_part'] / 2.0
            if not is_roll:
                detail['time'] += '<div style="color:gray;font-size:0.8rem">(备注：系统判定该休假包含年假%s天，事假%s天)</div>' % (annual_part, off_part)
                if job_record['cur_path_id']:
                    detail['mark'] = '<a href="detail_of_work_off.html?job_id=%s" target="_blank">跳转到销假条</a>' % job_record['cur_path_id']
            else:
                detail['time'] += '<div style="color:gray;font-size:0.8rem">(备注：系统判定该销假撤回年假%s天，事假%s天)</div>' % (-annual_part, -off_part)
                detail['mark'] = '<a href="detail_of_work_off.html?job_id=%s" target="_blank">跳转到休假条</a>' % job_record['cur_path_id']
        seq_field_map = {
            type_define.job_sequence_add: 'reason',
            type_define.job_sequence_pre_judge: 'pre_judgement',
            type_define.job_sequence_leader_judge: 'department_leader_judgement',
            type_define.job_sequence_via_leader_judge: 'via_leader_judgement',
            type_define.job_sequence_hr_leader_judge: 'hr_leader_judgement',
            type_define.job_sequence_main_leader_judge: 'main_leader_judgement',
            type_define.job_sequence_hr_record: 'hr_department_record',
        }
        invoker = job_record['invoker']
        next_sequence = job_record['cur_path_index']
        for node in job_nodes:
            sequence = node['branch_id']
            if sequence == type_define.job_sequence_add:
                detail['invoker'] = node['sender']
                detail['department'] = node['dept']
                if node['has_attachment']:
                    detail['attachment'] = yield self.job_dao.query_node_attachment_list(node['id'])
            content = self.unwrap_content(node['content'])
            content = '<div>%s</div><div style="color:gray;font-size:0.8rem">%s　　%s</div>' % (content, node['sender'], node['time'])
            detail[seq_field_map[sequence]] = content
        if next_sequence == type_define.job_sequence_pre_judge:
            detail['reply_desc'] = '已核'
            detail['reject_desc'] = '退回'
        elif next_sequence == type_define.job_sequence_hr_record:
            detail['reply_desc'] = '已备案'
            detail['reject_desc'] = '退回'
        else:
            next_next_sequence = job_sequence_map[job_record['type']][next_sequence]
            if next_next_sequence == type_define.job_sequence_hr_record:
                detail['reply_desc'] = '同意'
        detail['next_sequence'] = next_sequence
        detail['process_desc'] = seq_desc_map[next_sequence]
        status = job_record['status']
        mark = yield self.job_dao.query_job_mark(job_id, self.account_info['id'])
        detail['waiting'] = status == type_define.STATUS_JOB_PROCESSING and mark and mark['status'] == type_define.STATUS_JOB_MARK_WAITING 
        detail['status'] = job_status_desc[status]
        myself = invoker == self.account_info['id']
        detail['can_cancel'] = myself and next_sequence != type_define.job_sequence_hr_record and status == type_define.STATUS_JOB_PROCESSING
        detail['can_export'] = type_define.STATUS_JOB_COMPLETED == status
        detail['can_roll_back'] = myself and type_define.STATUS_JOB_COMPLETED == status and not is_roll and not job_record['cur_path_id']
        self.render('detail_of_work_off.html', account_info=self.account_info, detail=detail)

    def is_roll(self, job_record):
        return job_record['sub_type'] == type_define.TYPE_JOB_ROLL_BACK_LEAVE

    def unwrap_content(self, content):
        return content[1:-1]

    def get_leave_time_desc(self, leave_detail):
        day = leave_detail['half_day']/2.0
        if day < 0:
            day = -day
        return '自%s  至  %s，假期内工作日共%s天' % (leave_detail['begin_time'], leave_detail['end_time'], day)

    @gen.coroutine
    def get_annual(self, uid):
        annual = yield self.job_dao.query_user_annual_leave(uid)
        result = {
            'total': 0,
            'used': 0,
            'rest': 0,
        }
        if annual:
            result = {
                'total': annual['total'] / 2.0,
                'used': annual['used'] / 2.0,
            }
            result['rest'] = result['total'] - result['used']
        raise gen.Return(result)

    @gen.coroutine
    def post(self, *args, **kwargs):
        st = yield self.verify_user()
        if not st:
            return
        job_id = int(self.get_argument_and_check_it('job_id'))
        file_path = self.get_res_file_path('leave_table_%s.docx' % job_id, 'res/download/job_export', True)
        net_path = self.get_res_file_path('leave_table_%s.docx' % job_id, 'res/download/job_export')
        if not os.path.isfile(file_path):
            yield self.export_detail(job_id, file_path)
        else:
            yield self.export_detail(job_id, file_path)
            pass
        self.write_data({'path': net_path})


    @gen.coroutine
    def export_detail(self, job_id, file_path):
        job_record = yield self.job_dao.query_job_base_info(job_id)
        job_nodes = yield self.job_dao.query_job_node_list(job_id)
        leave_detail = yield self.job_dao.query_new_leave_detail(job_id)
        invoker = job_record['invoker']
        annual = yield self.get_annual(invoker)
        doc = Document()
        desc = u'请假' if not self.is_roll(job_record) else u'销假'
        title = doc.add_paragraph(u'深圳市东部传媒股份有限公司%s条' % desc)
        title.runs[0].font.size = Pt(24)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_map = {}
        self.create_table(doc, cell_map, desc)
        cell_map['type'].text = leave_detail['leave_type'].decode('utf8')
        cell_map['time'].text = u'\n%s\n' % self.get_leave_time_desc(leave_detail).decode('utf8')
        cell_map['total'].text = u'应休年假%s天' % annual['total']
        cell_map['used'].text = u'已休年假%s天' % annual['used']
        for node in job_nodes:
            sequence = node['branch_id']
            cell = None
            if sequence == type_define.job_sequence_add:
                cell_map['name'].text = node['sender'].decode('utf8')
                cell_map['dept'].text = node['dept'].decode('utf8')
                self.set_cell_content(cell_map['reason'], node)
                continue
            elif sequence == type_define.job_sequence_pre_judge:
                cell = cell_map['pre_judge']
            elif sequence == type_define.job_sequence_leader_judge:
                cell = cell_map['leader_judge']
            elif sequence == type_define.job_sequence_via_leader_judge:
                cell = cell_map['via_judge']
            elif sequence == type_define.job_sequence_main_leader_judge:
                cell = cell_map['main_judge']
            elif sequence == type_define.job_sequence_hr_record:
                cell = cell_map['hr_record']
            if cell is not None:
                self.set_cell_content(cell, node)
        sign = doc.add_paragraph(u'本人确认签名：')
        sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign.add_run(u'\t\t\t').underline = True
        doc.save(file_path)

    def create_table(self, doc, cell_map, desc):
        table = doc.add_table(rows=0, cols=4)
        table.style = 'TableGrid'
        table.autofit = True
        table.allow_autofit = True
        table.columns[0].width = Inches(1.5)

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'\n姓名\n')
        cells[2].text = u'部门'
        cell_map['name'] = cells[1]
        cell_map['dept'] = cells[3]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'\n%s类别\n' % desc)
        cell_map['type'] = cells[1]
        cell_map['total'] = cells[2]
        cell_map['used'] = cells[3]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'人事部初核意见')
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['pre_judge'] = cells[1]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'%s事由' % desc)
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['reason'] = cells[1]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'%s时间' % desc)
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['time'] = cells[1]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'部门负责人意见')
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['leader_judge'] = cells[1]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'\n分管领导意见\n')
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['via_judge'] = cells[1]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'\n主要负责人意见\n')
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['main_judge'] = cells[1]

        cells = table.add_row().cells
        self.set_field_cell(cells[0], u'人事部备案情况')
        cells[1].merge(cells[2])
        cells[1].merge(cells[3])
        cell_map['hr_record'] = cells[1]

        self.normalize_cells(table)

    def set_field_cell(self, cell, text):
        p = cell.paragraphs[0]
        p.add_run(text).bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def normalize_cells(self, table):
        columns = table.columns
        for col in columns:
            for cell in col.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def set_cell_content(self, cell, node):
        content = self.unwrap_content(node['content'])
        p = cell.add_paragraph(content.decode('utf8'))
        p = cell.add_paragraph(('%s    %s' % (node['sender'], node['time'].strftime('%Y年%m月%d日'))).decode('utf8'))
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.line_spacing = Pt(15)
        cell.add_paragraph()

